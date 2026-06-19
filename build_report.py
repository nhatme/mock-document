# -*- coding: utf-8 -*-
import sys, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DIAGRAMS = '/tmp/mock-document/diagrams'
OUT_PATH = '/home/nhatpx/practice/UIT-LAB/backup/financial-advisor/bao-cao-financial-advisor.docx'

# page_map: bookmark_name -> page number, filled in on the second pass.
PAGE_MAP_FILE = '/tmp/mock-document/toc_page_map.json'
try:
    with open(PAGE_MAP_FILE) as f:
        PAGE_MAP = json.load(f)
except FileNotFoundError:
    PAGE_MAP = {}

doc = Document()
toc_entries = []  # (level, text, bookmark_name)
_bookmark_counter = [0]


def add_bookmark(paragraph, name):
    _bookmark_counter[0] += 1
    bid = str(_bookmark_counter[0])
    p = paragraph._p
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), bid)
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bid)
    pPr = p.find(qn('w:pPr'))
    p.insert(1 if pPr is not None else 0, start)
    p.append(end)


def set_page_border(section, color='1F4E79', sz=24, val='single', space=24):
    sectPr = section._sectPr
    pgBorders = sectPr.find(qn('w:pgBorders'))
    if pgBorders is None:
        pgBorders = OxmlElement('w:pgBorders')
        sectPr.append(pgBorders)
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), str(space))
        el.set(qn('w:color'), color)
        pgBorders.append(el)


# Default style
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
rpr = normal.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = rpr.makeelement(qn('w:rFonts'), {})
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def set_cell_borders_none(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins_zero(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), '0')
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def add_title_page():
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo_w, text_w = Inches(0.8), Inches(5.5)
    table.columns[0].width = logo_w
    table.columns[1].width = text_w
    logo_cell, text_cell = table.rows[0].cells
    logo_cell.width = logo_w
    text_cell.width = text_w
    for cell in (logo_cell, text_cell):
        set_cell_borders_none(cell)
        set_cell_margins_zero(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.add_run().add_picture(f'{DIAGRAMS}/uit_logo.png', width=Inches(0.75))

    text_p = text_cell.paragraphs[0]
    text_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text_p.paragraph_format.left_indent = Inches(0.15)
    r = text_p.add_run('ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH\nTRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN')
    r.bold = True
    r.font.size = Pt(13)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('BÁO CÁO ĐỒ ÁN MÔN HỌC')
    r.bold = True
    r.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Môn: KỸ THUẬT LẬP TRÌNH TRÍ TUỆ NHÂN TẠO\n(AI ENGINEERING)')
    r.font.size = Pt(13)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('AI Tư Vấn Sản Phẩm Tài Chính Cho Người Trẻ Việt Nam')
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Hệ thống So khớp & Gợi ý Sản phẩm Tài chính có Giải thích')
    r.italic = True
    r.font.size = Pt(13)

    doc.add_paragraph()
    info_rows = [
        ('Chương trình Đào tạo', 'Liên thông Đại học Ngành Trí tuệ Nhân tạo'),
        ('Giảng viên hướng dẫn', 'TS. Đặng Văn Thìn'),
        ('Thành viên nhóm', '1. Phan Văn Nhật — 26410083'),
        ('', '2. Trương Quốc Bảo — 26410008'),
    ]
    info_table = doc.add_table(rows=0, cols=2)
    info_table.autofit = False
    itblPr = info_table._tbl.tblPr
    ilayout = OxmlElement('w:tblLayout')
    ilayout.set(qn('w:type'), 'fixed')
    itblPr.append(ilayout)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    label_w, value_w = Inches(1.9), Inches(3.0)
    for label, value in info_rows:
        row = info_table.add_row()
        row.cells[0].width = label_w
        row.cells[1].width = value_w
        for cell in row.cells:
            set_cell_borders_none(cell)
            set_cell_margins_zero(cell)
        lp = row.cells[0].paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if label:
            lr = lp.add_run(label + '  :')
            lr.font.size = Pt(12)
        vp = row.cells[1].paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        vp.paragraph_format.left_indent = Inches(0.15)
        vr = vp.add_run(value)
        vr.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(190)
    r = p.add_run('TP.HCM - 2026')
    r.bold = True

    set_page_border(doc.sections[0], color='1F4E79', sz=24, val='single', space=24)
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    pgBorders = new_section._sectPr.find(qn('w:pgBorders'))
    if pgBorders is not None:
        new_section._sectPr.remove(pgBorders)


def add_internal_hyperlink(paragraph, text, bookmark_name):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1F4E79')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def h1(text):
    bm = f'sec_{len(toc_entries)}'
    p = doc.add_heading(text, level=1)
    add_bookmark(p, bm)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    toc_entries.append((1, text, bm))
    return p


def h2(text):
    bm = f'sec_{len(toc_entries)}'
    p = doc.add_heading(text, level=2)
    add_bookmark(p, bm)
    toc_entries.append((2, text, bm))
    return p


def body(text):
    return doc.add_paragraph(text)


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    return p


def bold_lead(label, rest):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(label)
    r.bold = True
    p.add_run(rest)
    return p


def add_figure(path, caption, width=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = htext
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


add_title_page()

# --- placeholder paragraph for TOC; filled in after we know section order ---
toc_anchor_index = len(doc.paragraphs)
toc_heading = doc.add_paragraph()
toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = toc_heading.add_run('MỤC LỤC')
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
toc_body_para = doc.add_paragraph()  # filled later, kept as anchor for insert position
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Nội dung báo cáo đồ án')
r.bold = True
r.font.size = Pt(13)
doc.add_paragraph()

# 1. Discover
h1('1. Discover — Bài toán & Người dùng')
bold_lead('Vấn đề: ', 'người trẻ Việt Nam (sinh viên, mới đi làm) thường thiếu kiến thức tài chính, '
          'khó chọn đúng sản phẩm tiết kiệm, quỹ mở hay thẻ tín dụng phù hợp giữa hàng chục lựa chọn '
          'trên thị trường, và dễ nhận tư vấn thiên lệch hoặc chấp nhận rủi ro vượt quá khả năng.')
bold_lead('Người dùng: ', 'người trẻ 20–35 tuổi, thu nhập khởi điểm, ít kinh nghiệm đầu tư, quen dùng '
          'ứng dụng web/mobile, cần lời gợi ý rõ ràng và dễ hiểu hơn là báo cáo tài chính phức tạp.')
bold_lead('Mục tiêu sản phẩm: ', 'so khớp hồ sơ người dùng (tuổi, thu nhập, mục tiêu tài chính, khẩu vị '
          'rủi ro, kỳ hạn mong muốn) với danh mục sản phẩm tài chính (tiết kiệm, quỹ mở, thẻ tín dụng) → '
          'trả về danh sách gợi ý có thứ hạng, kèm lý do giải thích minh bạch theo từng tiêu chí và cảnh '
          'báo tuân thủ (compliance).')
bold_lead('Định vị (AI Engineering): ', 'đây KHÔNG phải chatbot hỏi-đáp tự do, mà là một engine so khớp '
          'sản phẩm có điểm số giải thích được (explainable scoring), xây trên hạ tầng truy xuất kết hợp '
          '(SQL + RAG) và LLM pluggable — không huấn luyện mô hình từ đầu, đúng tinh thần môn học '
          '(Chip Huyen, AI Engineering).')

add_figure(f'{DIAGRAMS}/lifecycle.png', 'Hình 1: Vòng đời dự án AI của Financial Advisor')

# 2. Design
h1('2. Design — Kiến trúc & Công nghệ')
body('Người dùng nhập hồ sơ tài chính trên giao diện web. Pipeline đa tác tử (multi-agent) xử lý tuần tự '
     'theo một đồ thị có hướng cố định (fixed DAG), từ làm giàu hồ sơ, truy xuất ứng viên hợp lệ, chấm '
     'điểm xếp hạng, đến kiểm tra tuân thủ, trước khi trả kết quả kèm lý do giải thích.')

h2('2.1 Pipeline tác tử (LangGraph DAG)')
body('UserProfile → [Profiler] → [Researcher] → [Recommender] → [Compliance] → RecommendResponse')

add_figure(f'{DIAGRAMS}/architecture.png', 'Hình 2: Kiến trúc hệ thống Financial Advisor')

h2('2.2 Stack công nghệ')
add_table(
    ['Tầng', 'Công nghệ'],
    [
        ['Frontend', 'Next.js 14, TypeScript, Tailwind CSS'],
        ['Backend', 'FastAPI + Uvicorn, Python 3.11'],
        ['Điều phối tác tử', 'LangGraph (StateGraph, fixed DAG)'],
        ['CSDL có cấu trúc', 'SQLite (dev) / PostgreSQL (prod), SQLAlchemy 2, Alembic'],
        ['Vector store', 'ChromaDB (local persistent)'],
        ['Embedding', 'intfloat/multilingual-e5-small (sentence-transformers, CPU)'],
        ['LLM', 'Pluggable: Anthropic → OpenAI → Ollama → Stub (tự dò)'],
        ['Logging', 'structlog (JSON có cấu trúc)'],
    ],
)

h2('2.3 Lý do chọn stack')
bullet('SQL (SQLite/PostgreSQL) lưu thuộc tính sản phẩm chính xác (lãi suất, điều kiện, hạn mức) — '
       'phù hợp cho lọc cứng theo ràng buộc định lượng.')
bullet('ChromaDB + embedding đa ngữ E5 cho RAG trên tài liệu điều khoản (T&C) tiếng Việt.')
bullet('LLM pluggable với chế độ Stub cho phép pipeline chạy đầy đủ ngay cả khi không có API key — '
       'dễ test, dễ demo, đúng tinh thần đồ án.')

# 3. Develop
h1('3. Develop — Pipeline & Kỹ thuật')

h2('3.1 Dữ liệu & Catalog')
bullet('Nguồn: catalog.json — 15 sản phẩm tài chính (tiết kiệm, quỹ mở, thẻ tín dụng).')
bullet('Mỗi sản phẩm có 1 file T&C dạng Markdown tương ứng (15 file) dùng để index vào ChromaDB.')
bullet('Schema: id, name, provider, type, risk_level, interest_rate_pct, expected_return_pct, '
       'annual_fee_vnd, min_amount_vnd, min_income_vnd, min_age, term_months, description, source_url.')
bullet('Seed idempotent qua scripts/seed.py (db.merge() cho SQL, vs.reset() cho vector index).')

h2('3.2 Truy xuất lai (Hybrid Retrieval)')
bold_lead('Bước 1 — Lọc SQL cứng: ', 'ProductRepository.filter_eligible(tuổi, thu nhập, loại sản phẩm) '
          'loại bỏ ngay các sản phẩm không đủ điều kiện trước khi đưa vào xếp hạng.')
bold_lead('Bước 2 — RAG theo từng ứng viên: ', 'với mỗi sản phẩm còn lại, ChromaDB tìm 2 đoạn liên quan '
          'nhất từ tài liệu T&C, dùng truy vấn tiếng Việt dựng từ goal + risk_appetite, được giới hạn '
          '(scoped) theo where={"product_id": p.id} để không lẫn nguồn trích dẫn giữa các sản phẩm.')
body('Embedding dùng quy ước tiền tố E5 (query: ... / passage: ...) để cosine similarity tính đúng.')

add_figure(f'{DIAGRAMS}/retrieval.png', 'Hình 3: Hybrid Retrieval — seed ngoại tuyến + truy vấn trực tuyến')

h2('3.3 Chấm điểm & Đa tác tử (Recommender)')
body('Recommender chấm điểm mỗi sản phẩm trên 5 chiều, sau đó sắp xếp giảm dần:')
add_table(
    ['Chiều', 'Ý nghĩa'],
    [
        ['risk_alignment', '1 − 0.4 × |risk người dùng − risk sản phẩm| (thang thứ tự 1–3)'],
        ['goal_alignment', 'Bảng tra (mục tiêu, loại sản phẩm) → trọng số 0.1–1.0'],
        ['income_compat', 'Thu nhập tháng ≥ thu nhập tối thiểu yêu cầu'],
        ['amount_feasibility', '20% thu nhập có đủ trang trải số tiền tối thiểu'],
        ['horizon_match', 'min(kỳ hạn mong muốn, kỳ hạn sản phẩm) / max(...)'],
    ],
)
bullet('Diversity cap: tối đa 2 sản phẩm/nhà cung cấp trong top-5, tránh độc quyền một bên.')
bullet('LLM chỉ được gọi tại Recommender — sinh lý giải 2 câu tiếng Việt (≤180 token) cho mỗi sản phẩm '
       'trong top-k. Profiler, Researcher, Compliance hoàn toàn rule-based, không gọi LLM.')

h2('3.4 Kiểm tra tuân thủ (Compliance)')
bullet('Loại bỏ sản phẩm risk_level=high khỏi danh sách nếu người dùng có risk_appetite=low.')
bullet('Cảnh báo nếu toàn bộ sản phẩm còn lại thuộc cùng một nhà cung cấp.')
bullet('Luôn đính kèm disclaimer bắt buộc ở cuối phản hồi.')
bullet('PII policy: hồ sơ người dùng không được log hay lưu trữ — mỗi request là stateless.')

# 4. Deliver
h1('4. Deliver — Demo')
body('Giao diện Next.js gồm form nhập hồ sơ (tuổi, thu nhập, mục tiêu, khẩu vị rủi ro, kỳ hạn, ghi chú), '
     'bảng so sánh top-3 sản phẩm, các product card (điểm số, ưu/nhược điểm, lý giải LLM), trace panel '
     'hiển thị quá trình xử lý của từng tác tử và cảnh báo tuân thủ, cùng một report modal toàn màn hình '
     'hỗ trợ in/lưu PDF.')
body('Kịch bản trình bày:')
bullet('1. Người dùng nhập hồ sơ: 24 tuổi, thu nhập 15.000.000đ/tháng, mục tiêu tiết kiệm, rủi ro thấp, '
       'kỳ hạn 12 tháng.')
bullet('2. Hệ thống trả về danh sách sản phẩm tiết kiệm/quỹ mở phù hợp, mỗi sản phẩm kèm điểm số theo '
       '5 chiều và lý giải ngắn bằng tiếng Việt.')
bullet('3. Xem trace panel để thấy rõ từng bước: Profiler → Researcher (trích dẫn nguồn T&C) → '
       'Recommender (điểm số) → Compliance (lọc rủi ro, cảnh báo).')
bullet('4. Mở report modal, in/lưu báo cáo PDF đầy đủ.')

# 5. Evaluate
h1('5. Evaluate — Đánh giá')
bullet('Bộ test pytest (backend/tests/test_pipeline.py) chạy với stub LLM, DB tạm, không gọi mạng — '
       'kiểm tra toàn bộ pipeline từ profile đến recommendation.')
bullet('Đánh giá thủ công trên các hồ sơ mẫu theo từng tổ hợp goal × risk_appetite để kiểm tra: '
       'thứ hạng hợp lý, diversity cap hoạt động đúng, compliance filter loại đúng sản phẩm rủi ro cao.')
bullet('Đánh giá RAG: trích dẫn (citation) trả về có khớp đúng với product_id đang xét không (do scoped '
       'theo product_id nên không bị lẫn nguồn giữa các sản phẩm khác nhau).')
bullet('Điểm cộng — trách nhiệm AI: disclaimer bắt buộc, PII không lưu trữ, lý giải minh bạch theo từng '
       'tiêu chí chấm điểm thay vì hộp đen.')

# 6. Difficulties & future
h1('6. Khó khăn & Hướng phát triển')
bold_lead('Khó khăn: ', 'cân bằng giữa chấm điểm rule-based và lý giải tự nhiên của LLM để vẫn giữ được '
          'tính giải thích được (explainability); catalog sản phẩm hiện còn nhỏ và mô phỏng; chưa có '
          'vòng phản hồi thực tế từ người dùng để hiệu chỉnh trọng số scoring.')
bold_lead('Hướng phát triển: ', 'mở rộng catalog sản phẩm thật, thêm vòng phản hồi người dùng (feedback '
          'loop) để tinh chỉnh trọng số 5 chiều, hỗ trợ đối thoại nhiều lượt quanh gợi ý, và triển khai '
          'production với PostgreSQL + GPU cho embedding.')

# 7. References
h1('7. Tài liệu tham khảo')
h2('Kỹ thuật RAG & LLM')
refs1 = [
    'P. Lewis et al. "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks." NeurIPS, 2020.',
    'J. Wei et al. "Chain-of-Thought Prompting Elicits Reasoning in LLMs." NeurIPS, 2022.',
    'C. Huyen. AI Engineering: Building Applications with Foundation Models. O’Reilly, 2025.',
    'Các kỹ thuật lập trình trong TTNT — slide RAG, Prompt Engineering, Agentic AI (TS. Đặng Văn Thìn, UIT, 2026).',
]
for i, r in enumerate(refs1, 1):
    body(f'{i}. {r}')

h2('Công cụ & framework')
refs2 = [
    'LangGraph / LangChain (langchain.com).',
    'ChromaDB vector database. trychroma.com.',
    'sentence-transformers — intfloat/multilingual-e5-small. huggingface.co.',
    'FastAPI (fastapi.tiangolo.com); SQLAlchemy 2 (sqlalchemy.org); Alembic.',
    'Next.js 14 (nextjs.org); Tailwind CSS (tailwindcss.com).',
    'Anthropic Claude API / OpenAI API / Ollama — pluggable LLM providers.',
]
for i, r in enumerate(refs2, 1):
    body(f'{i}. {r}')

# --- now build the actual TOC content into toc_body_para ---
tab_stops = toc_body_para.paragraph_format.tab_stops
tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

first = True
for level, text, bm in toc_entries:
    target = toc_body_para if first else doc.add_paragraph()
    first = False
    if target is not toc_body_para:
        tp = target.paragraph_format.tab_stops
        tp.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    if level == 2:
        target.paragraph_format.left_indent = Inches(0.3)
    run_label = target.add_run('')
    add_internal_hyperlink(target, text, bm)
    tab_run = target.add_run('\t' + str(PAGE_MAP.get(bm, '')))
    if level == 1:
        for r in target.runs:
            r.bold = True

# Move the freshly built TOC paragraphs to sit right after the "MỤC LỤC" heading
# (they were appended at the end of the body, so relocate them).
toc_heading_el = toc_heading._p
insert_after = toc_heading_el

# The first TOC entry's paragraph (toc_body_para) sits right after the heading already.
# The remaining entries were appended at the very end of the document via add_paragraph();
# walk back from the end to collect exactly that many paragraphs and relocate them.
extra_count = len(toc_entries) - 1
tail_paragraphs = doc.paragraphs[-extra_count:] if extra_count > 0 else []

for p in tail_paragraphs:
    p._p.getparent().remove(p._p)
    insert_after.addnext(p._p)
    insert_after = p._p

# finally move toc_body_para right after heading (before the tail ones we just placed,
# so overall order is heading -> toc_body_para -> tail_paragraphs)
toc_body_para._p.getparent().remove(toc_body_para._p)
toc_heading_el.addnext(toc_body_para._p)

doc.save(OUT_PATH)
print('saved')
